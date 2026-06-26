from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "cv" / "CV-Chinese-AI-Security.md"
OUTPUT = Path(r"D:\workspace\cv for work\CV Chinese AI Security.docx")


def set_east_asian_font(run, name: str = "Microsoft YaHei") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", name)


def build_docx() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "Microsoft YaHei")
    style.font.size = Pt(9.5)

    for raw in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue

        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(18)
            set_east_asian_font(run)
        elif line.startswith("## "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(line[3:])
            run.bold = True
            run.font.size = Pt(12)
            set_east_asian_font(run)
        elif line.startswith("### "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(10.5)
            set_east_asian_font(run)
        elif line.startswith("- "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            paragraph.paragraph_format.space_after = Pt(1)
            bullet = paragraph.add_run("- ")
            bullet.bold = True
            set_east_asian_font(bullet)
            run = paragraph.add_run(line[2:])
            set_east_asian_font(run)
        else:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)
            if "浙江工商大学 |" in line:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line)
            set_east_asian_font(run)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()

